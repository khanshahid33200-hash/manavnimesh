import subprocess
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

def create_html_resume():
    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Resume - MANAV NIMESH</title>
<style>
  @page {
    size: letter;
    margin: 0.35in 0.55in;
  }
  
  * {
    box-sizing: border-box;
    margin: 0;
    padding: 0;
  }

  body {
    font-family: 'Times New Roman', Times, serif;
    color: #111111;
    font-size: 9.5pt;
    line-height: 1.25;
    background-color: #ffffff;
  }

  .container {
    width: 100%;
  }

  /* Header Styling */
  .header {
    text-align: center;
    margin-bottom: 4px;
  }

  .name {
    font-family: 'Times New Roman', Times, serif;
    font-size: 19pt;
    font-weight: bold;
    letter-spacing: 1.5px;
    text-transform: uppercase;
    color: #000000;
    margin-bottom: 3px;
  }

  .contact-info {
    font-size: 9pt;
    color: #222222;
    display: flex;
    justify-content: center;
    align-items: center;
    gap: 8px;
  }

  .email-text {
    font-size: 8.5pt;
    text-transform: lowercase;
  }

  .divider {
    border: none;
    border-top: 1px solid #666666;
    margin: 5px 0 6px 0;
  }

  /* Section Styling */
  .section {
    margin-bottom: 4px;
  }

  .section-title {
    font-family: Arial, Helvetica, sans-serif;
    font-size: 10pt;
    font-weight: bold;
    text-transform: uppercase;
    color: #000000;
    letter-spacing: 0.5px;
    margin-bottom: 3px;
  }

  .text-block {
    text-align: justify;
    font-size: 9.5pt;
    color: #222222;
    line-height: 1.3;
  }

  /* 2-Column Grid Layout */
  .two-column-grid {
    display: flex;
    justify-content: space-between;
    gap: 20px;
  }

  .col {
    flex: 1;
  }

  /* Lists Styling */
  ul.bullet-list {
    list-style-type: disc;
    margin-left: 16px;
    margin-top: 2px;
    margin-bottom: 3px;
  }

  li {
    font-size: 9.3pt;
    margin-bottom: 1.5px;
    color: #222222;
  }

  /* Education Section */
  .edu-item {
    margin-bottom: 3px;
  }

  .edu-item:last-child {
    margin-bottom: 0;
  }

  .edu-degree {
    font-weight: bold;
    font-size: 9.8pt;
    color: #000000;
  }

  .edu-school {
    font-style: italic;
    font-size: 9.2pt;
    color: #333333;
  }
</style>
</head>
<body>

<div class="container">
  <!-- Header -->
  <div class="header">
    <div class="name">MANAV NIMESH</div>
    <div class="contact-info">
      <span>📍 Jaipur, India</span>
      <span>📞 8739997065</span>
      <span>|</span>
      <span>✉️ <span class="email-text">manavnimesh12@gmail.com</span></span>
    </div>
  </div>

  <hr class="divider" />

  <!-- Professional Summary -->
  <div class="section">
    <div class="section-title">PROFESSIONAL SUMMARY</div>
    <p class="text-block">
      Dedicated and student-focused Science Teacher & Private Tutor with a strong academic background in Physics, Chemistry, and Mathematics. Passionate about creating engaging classroom experiences, 1-on-1 tutoring, and helping students build strong conceptual understanding. Skilled in classroom management, lesson planning, and effective communication. Seeking a Science Teacher / Private Tutor position where academic knowledge and teaching skills contribute to student success.
    </p>
  </div>

  <hr class="divider" />

  <!-- 2 Column Side-by-Side Skills Section -->
  <div class="two-column-grid">
    <!-- Column 1 -->
    <div class="col">
      <div class="section">
        <div class="section-title">TEACHING & TUTORING SKILLS</div>
        <ul class="bullet-list">
          <li>Private Tutoring & 1-on-1 Mentoring</li>
          <li>Science Teaching (Physics, Chemistry, Math)</li>
          <li>Lesson Planning & Curriculum Design</li>
          <li>Classroom & Batch Management</li>
          <li>Student Assessment & Mock Tests</li>
          <li>Activity Based Learning</li>
          <li>Doubt Solving & Concept Building</li>
        </ul>
      </div>

      <div class="section">
        <div class="section-title">TECHNICAL SKILLS</div>
        <ul class="bullet-list">
          <li>MS Word, MS Excel, MS PowerPoint</li>
          <li>Basic Computer Knowledge</li>
          <li>Data Entry</li>
        </ul>
      </div>
    </div>

    <!-- Column 2 -->
    <div class="col">
      <div class="section">
        <div class="section-title">SOFT SKILLS</div>
        <ul class="bullet-list">
          <li>Creative Problem Solving</li>
          <li>Time Management</li>
          <li>Interpersonal Skills</li>
          <li>Verbal and Written Communication</li>
          <li>Teamwork</li>
          <li>Leadership</li>
        </ul>
      </div>

      <div class="section">
        <div class="section-title">STRENGTHS</div>
        <ul class="bullet-list">
          <li>Strong subject knowledge in Science</li>
          <li>Positive classroom attitude</li>
          <li>Good presentation skills</li>
          <li>Quick learner</li>
          <li>Responsible and disciplined</li>
        </ul>
      </div>
    </div>
  </div>

  <hr class="divider" />

  <!-- Education -->
  <div class="section">
    <div class="section-title">EDUCATION</div>
    <div class="edu-item">
      <div class="edu-degree">Master of Science (M.Sc.), Chemistry</div>
      <div class="edu-school">University of Rajasthan, Jaipur — 2025</div>
    </div>
    <div class="edu-item">
      <div class="edu-degree">Bachelor of Science & Bachelor of Education (B.Sc. B.Ed.), PCM</div>
      <div class="edu-school">University of Rajasthan, Jaipur — 2023</div>
    </div>
    <div class="edu-item">
      <div class="edu-degree">Bachelor of Computer Applications (BCA)</div>
      <div class="edu-school">Vardhman Mahaveer Open University, Kota — Expected 2026</div>
    </div>
  </div>

  <hr class="divider" />

  <!-- Courses & Certifications -->
  <div class="section">
    <div class="section-title">COURSES & CERTIFICATIONS</div>
    <ul class="bullet-list">
      <li>Central Teacher Eligibility Test (CTET) — 2024</li>
      <li>Rajasthan Eligibility Examination for Teachers (REET Rajasthan) — 2025</li>
      <li>Rajasthan State Certificate in Information Technology (RSCIT) — 2021</li>
      <li>Information Practices, CBSE Senior Secondary Education — 2018</li>
    </ul>
  </div>

  <hr class="divider" />

  <!-- Career Objective -->
  <div class="section">
    <div class="section-title">CAREER OBJECTIVE</div>
    <p class="text-block">
      To work as a Science Teacher or Private Tutor in a reputed school or tutoring setup where I contribute to students' academic growth through effective teaching methods, practical learning, and continuous professional development.
    </p>
  </div>

  <hr class="divider" />

  <!-- Languages at the very bottom -->
  <div class="section">
    <div class="section-title">LANGUAGES</div>
    <ul class="bullet-list">
      <li>Hindi</li>
      <li>English</li>
    </ul>
  </div>

  <hr class="divider" />

</div>

</body>
</html>
"""
    with open("MANAV_NIMESH_RESUME.html", "w", encoding="utf-8") as f:
        f.write(html_content)
    print("Saved MANAV_NIMESH_RESUME.html")

def build_pdf():
    edge_path = r'C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe'
    html_path = os.path.abspath("MANAV_NIMESH_RESUME.html")
    pdf_path = os.path.abspath("MANAV_NIMESH_RESUME.pdf")

    cmd = [
        edge_path,
        '--headless=new',
        '--disable-gpu',
        f'--print-to-pdf={pdf_path}',
        '--no-pdf-header-footer',
        html_path
    ]
    res = subprocess.run(cmd, capture_output=True, text=True)
    print("PDF Conversion Return Code:", res.returncode)
    print("PDF exists:", os.path.exists(pdf_path))

def create_docx_resume():
    doc = docx.Document()
    
    # Page setup - Margins 0.35 in top/bottom, 0.5 in left/right
    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Base styling
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(9.5)
    normal_font.color.rgb = RGBColor(17, 17, 17)

    def add_bottom_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="3" w:color="666666"/></w:pBdr>')
        pPr.append(pBdr)

    # Name Header
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("MANAV NIMESH")
    run_name.bold = True
    run_name.font.size = Pt(19)
    run_name.font.name = 'Times New Roman'

    # Contact Info
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(4)
    run_contact = p_contact.add_run("📍 Jaipur, India   📞 8739997065   |   ✉️ ")
    run_contact.font.size = Pt(9)
    run_contact.font.name = 'Times New Roman'
    
    run_email = p_contact.add_run("manavnimesh12@gmail.com")
    run_email.font.size = Pt(8.5)
    run_email.font.name = 'Times New Roman'
    add_bottom_border(p_contact)

    def add_section_header(title, p_obj=None):
        p = p_obj if p_obj else doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        return p

    def add_horizontal_line():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_bottom_border(p)

    # Professional Summary
    add_section_header("PROFESSIONAL SUMMARY")
    p_sum = doc.add_paragraph(
        "Dedicated and student-focused Science Teacher & Private Tutor with a strong academic background in Physics, Chemistry, and Mathematics. Passionate about creating engaging classroom experiences, 1-on-1 tutoring, and helping students build strong conceptual understanding. Skilled in classroom management, lesson planning, and effective communication. Seeking a Science Teacher / Private Tutor position where academic knowledge and teaching skills contribute to student success."
    )
    p_sum.paragraph_format.space_after = Pt(4)
    p_sum.paragraph_format.line_spacing = 1.15
    add_horizontal_line()

    # 2-Column Table for Skills & Strengths
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.7)
    table.columns[1].width = Inches(3.7)
    
    # Remove table borders
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(tblBorders)

    cell_left = table.rows[0].cells[0]
    cell_right = table.rows[0].cells[1]

    # Cell Left Content: TEACHING SKILLS & TECHNICAL SKILLS
    p_l1 = cell_left.paragraphs[0]
    add_section_header("TEACHING & TUTORING SKILLS", p_l1)
    teaching_skills = [
        "Private Tutoring & 1-on-1 Mentoring",
        "Science Teaching (Physics, Chemistry, Math)",
        "Lesson Planning & Curriculum Design",
        "Classroom & Batch Management",
        "Student Assessment & Mock Tests",
        "Activity Based Learning",
        "Doubt Solving & Concept Building"
    ]
    for skill in teaching_skills:
        p = cell_left.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(skill)
        r.font.size = Pt(9.3)

    p_l2 = cell_left.add_paragraph()
    add_section_header("TECHNICAL SKILLS", p_l2)
    tech_skills = [
        "MS Word, MS Excel, MS PowerPoint",
        "Basic Computer Knowledge",
        "Data Entry"
    ]
    for skill in tech_skills:
        p = cell_left.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(skill)
        r.font.size = Pt(9.3)

    # Cell Right Content: SOFT SKILLS & STRENGTHS
    p_r1 = cell_right.paragraphs[0]
    add_section_header("SOFT SKILLS", p_r1)
    soft_skills = [
        "Creative Problem Solving",
        "Time Management",
        "Interpersonal Skills",
        "Verbal and Written Communication",
        "Teamwork",
        "Leadership"
    ]
    for skill in soft_skills:
        p = cell_right.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(skill)
        r.font.size = Pt(9.3)

    p_r2 = cell_right.add_paragraph()
    add_section_header("STRENGTHS", p_r2)
    strengths = [
        "Strong subject knowledge in Science",
        "Positive classroom attitude",
        "Good presentation skills",
        "Quick learner",
        "Responsible and disciplined"
    ]
    for strg in strengths:
        p = cell_right.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(strg)
        r.font.size = Pt(9.3)

    add_horizontal_line()

    # Education
    add_section_header("EDUCATION")
    edu_list = [
        ("Master of Science (M.Sc.), Chemistry", "University of Rajasthan, Jaipur — 2025"),
        ("Bachelor of Science & Bachelor of Education (B.Sc. B.Ed.), PCM", "University of Rajasthan, Jaipur — 2023"),
        ("Bachelor of Computer Applications (BCA)", "Vardhman Mahaveer Open University, Kota — Expected 2026")
    ]
    for degree, school in edu_list:
        p_deg = doc.add_paragraph()
        p_deg.paragraph_format.space_before = Pt(2)
        p_deg.paragraph_format.space_after = Pt(0.5)
        r1 = p_deg.add_run(degree)
        r1.bold = True
        r1.font.size = Pt(9.8)

        p_sch = doc.add_paragraph()
        p_sch.paragraph_format.space_before = Pt(0)
        p_sch.paragraph_format.space_after = Pt(3)
        r2 = p_sch.add_run(school)
        r2.italic = True
        r2.font.size = Pt(9.2)
        r2.font.color.rgb = RGBColor(51, 51, 51)

    add_horizontal_line()

    # Certifications
    add_section_header("COURSES & CERTIFICATIONS")
    certs = [
        "Central Teacher Eligibility Test (CTET) — 2024",
        "Rajasthan Eligibility Examination for Teachers (REET Rajasthan) — 2025",
        "Rajasthan State Certificate in Information Technology (RSCIT) — 2021",
        "Information Practices, CBSE Senior Secondary Education — 2018"
    ]
    for cert in certs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(cert)
        r.font.size = Pt(9.3)

    add_horizontal_line()

    # Career Objective
    add_section_header("CAREER OBJECTIVE")
    p_obj = doc.add_paragraph(
        "To work as a Science Teacher or Private Tutor in a reputed school or tutoring setup where I contribute to students' academic growth through effective teaching methods, practical learning, and continuous professional development."
    )
    p_obj.paragraph_format.space_after = Pt(4)
    p_obj.paragraph_format.line_spacing = 1.15
    add_horizontal_line()

    # Languages at the very bottom
    add_section_header("LANGUAGES")
    langs = ["Hindi", "English"]
    for lang in langs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(lang)
        r.font.size = Pt(9.3)
    add_horizontal_line()

    doc.save("MANAV_NIMESH_RESUME.docx")
    print("Saved MANAV_NIMESH_RESUME.docx")

if __name__ == "__main__":
    create_html_resume()
    build_pdf()
    create_docx_resume()
